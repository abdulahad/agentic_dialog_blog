import streamlit as st
import markdown2
import io
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from bs4 import BeautifulSoup
from io import BytesIO
from src.services.corrector import Corrector
from src.services.meta_generator import MetaGenerator
from src.ui.correction import UIHandlerCorrection
from src.page_config import PageConfig
from src.testing.load_correction_examples import load_example, load_none

#### Setting up page config 
config = PageConfig(page_title='Correction')
prefix, language = config.configurate_page()

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._element.append(hyperlink)
    return hyperlink

def markdown_to_docx(markdown_text, brief, meta):
    html_text = markdown2.markdown(markdown_text, extras=["tables"])
    html_brief = markdown2.markdown(brief)
    html_meta = markdown2.markdown(meta)

    soup = BeautifulSoup(html_text, 'html.parser')
    soup_brief = BeautifulSoup(html_brief, 'html.parser')
    soup_meta = BeautifulSoup(html_meta, 'html.parser')

    doc = Document()

    def add_paragraph(element, style=None):
        p = doc.add_paragraph(style=style)
        for child in element.children:
            if child.name == 'a':
                add_hyperlink(p, child.get_text(), child['href'])
            elif child.name == 'strong':
                run = p.add_run(child.get_text())
                run.bold = True
            elif child.name:
                p.add_run(child.get_text())
            else:
                p.add_run(child)

    def add_list_item(element, list_style):
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style=list_style)
            for child in li.children:
                if child.name == 'a':
                    add_hyperlink(p, child.get_text(), child['href'])
                elif child.name == 'strong':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name:
                    p.add_run(child.get_text())
                else:
                    p.add_run(child)

    def add_table(element):
        rows = element.find_all('tr')
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                cell_text = ''.join(cell.stripped_strings)
                table.cell(i, j).text = cell_text

    tag_to_style = {
        'h1': 'Heading 1',
        'h2': 'Heading 2',
        'h3': 'Heading 3',
        'h4': 'Heading 4',
        'h5': 'Heading 5',
        'h6': 'Heading 6',
        'p': None,
        'ul': 'List Bullet',
        'ol': 'List Number',
        'table': 'Table'
    }

    def process_element(element):
        if element.name in tag_to_style:
            if element.name == 'table':
                add_table(element)
            elif element.name in ['ul', 'ol']:
                add_list_item(element, tag_to_style[element.name])
            else:
                add_paragraph(element, tag_to_style[element.name])

    def process_soup(soup):
        for element in soup.descendants:
            if element.name:
                process_element(element)

    process_soup(soup_brief)
    doc.add_page_break()
    process_soup(soup)
    doc.add_page_break()
    process_soup(soup_meta)

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f


#### Set up example (Load placeholder st.session_state variables to visualize and test UI)
# load_example()
# # Set language for example
# language = 'Francais'

#### Set up UI
ui_handler = UIHandlerCorrection(prefix=prefix, language=language, current_meta_stuff=st.session_state.get('current_meta_stuff', None))
ui_handler.display_ui()

user_input = ui_handler.get_user_input()
current_user_prompt = user_input['current_user_prompt']
generate_meta = user_input['generate_meta']

# NOTE: Perhaps we want to create an assistant here that has the context of the seo analysis in order 
# for it to be able to answer problems directly related to the text
# TODO: Create a button that allows the user to re-run SEO analysis on the current version of the text

if generate_meta:
    user_text = st.session_state['final_versions'].get_text(language)
    meta = MetaGenerator(article=user_text, keyword=st.session_state['theme'])
    meta_title = meta.generate_title()
    meta_desc = meta.generate_desc()
    st.session_state['current_meta_stuff'] = {'meta_title': meta_title, 'meta_description': meta_desc}
    st.rerun()

if current_user_prompt:
    with st.spinner("Correction en cours..."):
        # If user sends in a prompt, modify text with ChatGPT
        user_text = st.session_state['final_versions'].get_text(language)
        corrector = Corrector(text=user_text, prompt=current_user_prompt, thread=st.session_state.get('correction_thread', None)) 
        st.session_state['correction_thread'] = corrector.thread

        prompt_result = corrector.apply_prompt_to_text()
        st.session_state['final_versions'].update_text(language, prompt_result) 

        # TODO: Re-run a quick update on the SEO analysis (Update checklist + scoring maybe)
        new_seo_analysis = None
        st.session_state['final_versions'].update_seo_rating(language, new_seo_analysis) 

        # Re-run the page to ensure the updated text is displayed
        st.rerun()  

# Download document
if 'final_versions' in st.session_state:
    final_text = markdown_to_docx(st.session_state['text'], st.session_state['user_brief'], st.session_state['meta'])
    # Create the download button
    st.download_button(
        label="Télécharger le document",
        data=final_text,
        file_name='document.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

